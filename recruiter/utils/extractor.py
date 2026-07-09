import os
import io
import fitz  # PyMuPDF
fitz.TOOLS.mupdf_warnings(False)  # suppress MuPDF xref errors
import pytesseract
from PIL import Image
from docx import Document

# Set Tesseract Path for Windows/Linux
if os.name == 'nt': # Windows
    tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    if os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
else: # Linux
    pytesseract.pytesseract.tesseract_cmd = 'tesseract'

async def extract_text_from_file(file_path: str = None, file_content: bytes = None, filename: str = None) -> str:
    """
    Extracts text from a given file path or raw bytes.
    If file_content is provided, it uses that directly.
    """
    if file_path:
        _, ext = os.path.splitext(file_path)
    elif filename:
        _, ext = os.path.splitext(filename)
    else:
        raise ValueError("Either file_path or filename must be provided.")
        
    ext = ext.lower()
    
    if ext == ".pdf":
        text = await extract_text_from_pdf(file_path, file_content)

        if not text or len(text.strip()) < 10:
           print("Normal PDF extraction failed. Attempting OCR...")
           text = await extract_text_with_ocr(file_path, file_content)

        # FINAL VALIDATION AFTER OCR
        if not text or len(text.strip()) < 50:
           raise ValueError(
               "Invalid document. Unable to extract sufficient resume content."
          )
        return text
    elif ext in [".webp", ".png", ".jpg", ".jpeg"]:
        print(f"Image file detected ({ext}). Using OCR directly...")
        return await extract_image_text_with_ocr(file_path, file_content)
    elif ext == ".docx":
        return await extract_text_from_docx(file_path, file_content)
    elif ext == ".doc":
        raise ValueError(".doc format is not yet supported. Please use .docx or .pdf.")
    else:
        raise ValueError(f"Unsupported file format: {ext}")

async def extract_image_text_with_ocr(file_path: str = None, file_content: bytes = None) -> str:
    """Extracts text directly from an image file."""
    try:
        if file_content:
            img = Image.open(io.BytesIO(file_content))
        else:
            img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        print(f"Image OCR Error: {e}")
        raise ValueError(f"OCR failed for image: {str(e)}.")

async def extract_text_from_pdf(file_path: str = None, file_content: bytes = None) -> str:
    text = ""
    doc = None
    try:
        if file_content:
            doc = fitz.open(stream=file_content, filetype="pdf")
        else:
            doc = fitz.open(file_path)
            
        for page in doc:
            page_text = page.get_text("text") 
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"PyMuPDF error: {e}")
    finally:
        if doc:
            doc.close()
    return text.strip()

async def extract_text_with_ocr(file_path: str = None, file_content: bytes = None) -> str:
    """Converts PDF pages to images and uses Tesseract to extract text."""
    text = ""
    doc = None
    try:
        if file_content:
            doc = fitz.open(stream=file_content, filetype="pdf")
        else:
            doc = fitz.open(file_path)
        
        for i in range(len(doc)):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_text = pytesseract.image_to_string(img)
            if page_text:
                text += page_text + "\n"
        
    except Exception as e:
        print(f"OCR Error: {e}")
        raise ValueError(f"OCR failed: {str(e)}")
    finally:
        if doc:
            doc.close()
    return text.strip()

async def extract_text_from_docx(file_path: str = None, file_content: bytes = None) -> str:
    if file_content:
        doc = Document(io.BytesIO(file_content))
    else:
        doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    row_text.append(t)
            if row_text:
                parts.append(" | ".join(row_text))
    return "\n".join(parts)
