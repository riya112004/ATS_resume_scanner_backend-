import os
import logging
from PIL import Image
from recruiter.core.config import settings

logger = logging.getLogger("converter")

async def convert_to_pdf(file_path: str, filename: str):
    ext = os.path.splitext(filename)[1].lower()
    pdf_ext = ".pdf"
    pdf_path = os.path.splitext(file_path)[0] + pdf_ext
    pdf_filename = os.path.splitext(filename)[0] + pdf_ext

    if ext == pdf_ext:
        return file_path, filename

    if ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"):
        try:
            img = Image.open(file_path)
            if img.mode in ("RGBA", "LA", "P"):
                img = img.convert("RGB")
            img.save(pdf_path, "PDF", resolution=100.0)
            os.remove(file_path)
            logger.info(f"Converted image {filename} → {pdf_filename}")
            return pdf_path, pdf_filename
        except Exception as e:
            logger.error(f"Image → PDF conversion failed for {filename}: {e}")
            return file_path, filename

    if ext in (".docx", ".doc"):
        try:
            import subprocess
            import asyncio
            proc = await asyncio.create_subprocess_exec(
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", os.path.dirname(pdf_path),
                file_path,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.DEVNULL,
            )
            rc = await proc.wait()
            if rc == 0 and os.path.exists(pdf_path):
                os.remove(file_path)
                logger.info(f"Converted {filename} → {pdf_filename} via LibreOffice")
                return pdf_path, pdf_filename
            raise Exception(f"LibreOffice exit code {rc}")
        except Exception as e:
            logger.warning(f"LibreOffice failed for {filename}, trying fallback: {e}")
            try:
                from docx import Document
                from fpdf import FPDF
                doc = Document(file_path)
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Helvetica", "", 11)
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        safe = text.encode("latin-1", "replace").decode("latin-1")
                        pdf.multi_cell(0, 5, safe, new_x="LMARGIN", new_y="NEXT")
                    else:
                        pdf.ln(2)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            t = cell.text.strip()
                            if t:
                                safe = t.encode("latin-1", "replace").decode("latin-1")
                                pdf.multi_cell(0, 5, safe, new_x="LMARGIN", new_y="NEXT")
                pdf.output(pdf_path)
                if os.path.exists(pdf_path):
                    os.remove(file_path)
                    logger.info(f"Fallback DOCX → PDF done for {filename}")
                    return pdf_path, pdf_filename
            except Exception as e2:
                logger.error(f"DOCX → PDF fallback also failed for {filename}: {e2}")
            return file_path, filename

    return file_path, filename